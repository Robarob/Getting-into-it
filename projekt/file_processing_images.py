import utils
import fitz
from PIL import Image, ImageFilter
from pathlib import Path
from docx.shared import Cm
from pytesseract import image_to_string
from docx import Document
from file_processing_utils import (
    check_image_for_text,
    SUPPORTED_IMAGE_EXTS,
    TEMP_IMAGE_DIR,
    cleanup_temp_images,
)

# Kleinere Standardgröße (1000×1000 px) beschleunigt Verkleinerung und OCR
DEFAULT_MAX_SIZE = (1000, 1000)

def process_images_in_folder(
    folder_path: Path | str,
    output_docx_path: Path | str,
    *,
    max_size: tuple[int, int] = DEFAULT_MAX_SIZE,
    temp_dir: Path | None = None,
) -> None:
    """
    Verarbeitet alle Bilddateien in einem Ordner und schreibt das Ergebnis in ein Word‑Dokument.
    Für jedes Bild wird geprüft, ob Text vorhanden ist. Enthält das Bild Text,
    wird der erkannte Text im Dokument abgelegt. Andernfalls wird das farbige Bild eingefügt.
    """
    utils.debug_log(f"Starte Verarbeitung des Ordners: {folder_path}")

    try:
        input_files = [
            f
            for f in Path(folder_path).rglob("*")
            if f.suffix.lower() in SUPPORTED_IMAGE_EXTS
        ]
        utils.debug_log(f"Gefundene Bilddateien im Ordner: {len(input_files)}")

        word_doc = Document()

        for file_path in input_files:
            utils.debug_log(f"Starte Verarbeitung für Datei: {file_path}")
            try:
                result = extract_images_from_image(
                    file_path, max_size=max_size, temp_dir=temp_dir
                )
                utils.debug_log(f"Verarbeitungsergebnis für {file_path}: {result}")

                if result:
                    if (
                        "text" in result
                        and isinstance(result["text"], str)
                        and result["text"].strip()
                    ):
                        word_doc.add_paragraph(f"--- {file_path.name} ---")
                        word_doc.add_paragraph(result["text"])
                        utils.debug_log(f"Text aus Bild hinzugefügt: {file_path.name}")
                    else:
                        if "images" in result and result["images"]:
                            for image_path_str in result["images"]:
                                image_path = Path(image_path_str)
                                if image_path.exists():
                                    try:
                                        word_doc.add_paragraph(f"--- {file_path.name} ---")
                                        word_doc.add_picture(str(image_path), width=Cm(15))
                                        utils.debug_log(
                                            f"Bild ohne Text erfolgreich eingefügt: {image_path}"
                                        )
                                    except Exception as e:
                                        utils.log_error(
                                            f"Fehler beim Einfügen von Bild {image_path}: {e}"
                                        )
                                else:
                                    utils.log_error(
                                        f"Bilddatei existiert nicht: {image_path}"
                                    )
                else:
                    utils.debug_log(f"Verarbeitung fehlgeschlagen für Datei: {file_path}")
            except Exception as e:
                utils.log_error(f"Fehler bei der Verarbeitung von Datei {file_path}: {e}")

        word_doc.save(output_docx_path)
        utils.debug_log(f"Word-Dokument erfolgreich gespeichert: {output_docx_path}")

    except Exception as e:
        utils.log_error(f"Fehler während der Ordnerverarbeitung: {e}")

def extract_images_from_image(
    file_path: Path | str,
    *,
    max_size: tuple[int, int] = DEFAULT_MAX_SIZE,
    temp_dir: Path | None = None,
) -> dict:
    """
    Extrahiert Text und verarbeitet Bilder aus einer Bilddatei:
      1) Farbbild laden (WebP nach RGB konvertieren).
      2) Graustufenkopie für OCR erzeugen und optional schärfen.
      3) Beide Versionen verkleinern (max_size).
      4) Temporäre JPEGs speichern (komprimiert).
      5) Prüfen, ob Text enthalten ist. Wenn ja, OCR durchführen, sonst Bild zurückgeben.
    """
    try:
        utils.debug_log(f"Beginne Verarbeitung von Bild: {file_path}")
        tmp_dir = Path(temp_dir) if temp_dir else TEMP_IMAGE_DIR
        tmp_dir.mkdir(parents=True, exist_ok=True)

        with Image.open(file_path) as img_color:
            if img_color.format == "WEBP":
                img_color = img_color.convert("RGB")
                utils.debug_log(f"WebP-Bild konvertiert zu RGB: {file_path}")

            img_gray = img_color.convert("L")
            utils.debug_log(f"Bild zu Graustufen konvertiert: {file_path}")

            img_gray = img_gray.filter(ImageFilter.SHARPEN)
            utils.debug_log(f"Graustufen-Bild geschärft: {file_path}")

            img_color.thumbnail(max_size)
            img_gray.thumbnail(max_size)
            utils.debug_log(f"Bildgröße angepasst: {file_path}")

            temp_color_path = save_temp_image(
                img_color, file_path, suffix="_color", temp_dir=tmp_dir
            )
            temp_gray_path = save_temp_image(
                img_gray, file_path, suffix="_gray", temp_dir=tmp_dir
            )

            if not temp_color_path or not temp_gray_path:
                utils.log_error(
                    f"Temporäre Bilder konnten nicht gespeichert werden: {file_path}"
                )
                return {"text": "[Kein Text erkannt]", "images": []}

        utils.debug_log(
            f"Temporäre Dateien gespeichert: FARBE={temp_color_path}, GRAU={temp_gray_path}"
        )

        is_text_only = check_image_for_text(temp_gray_path)
        utils.debug_log(
            f"Textprüfung abgeschlossen: {'Text erkannt' if is_text_only else 'Kein Text erkannt'}"
        )

        if not is_text_only:
            utils.debug_log(
                f"Bild {file_path} enthält keinen Text. Gebe Farbbildpfad zurück."
            )
            return {"text": "", "images": [str(temp_color_path)]}

        # OCR
        try:
            with Image.open(temp_gray_path) as gray_for_ocr:
                ocr_text = utils.run_with_timeout(
                    image_to_string,
                    gray_for_ocr,
                    lang="deu+eng",
                    timeout=120,
                )
            if not ocr_text or not ocr_text.strip():
                utils.log_error(f"OCR-Ergebnis leer: {file_path}")
                ocr_text = "[Kein Text erkannt]"
            else:
                ocr_text = utils.sanitize_text_line(ocr_text)
                utils.debug_log(f"OCR-Text: {ocr_text[:50]}...")
            return {"text": ocr_text.strip(), "images": []}
        except Exception as e:
            utils.log_error(f"Fehler bei der OCR-Verarbeitung: {e}")
            return {"text": "[Kein Text erkannt]", "images": []}
    except Exception as e:
        utils.log_error(f"Fehler bei der Verarbeitung von {file_path}: {e}")
        return {"text": "[Kein Text erkannt]", "images": []}

def save_temp_image(
    image: Image.Image,
    original_path: Path | str,
    *,
    suffix: str = "",
    temp_dir: Path | None = None,
    format: str = "JPEG",
    quality: int = 80,
) -> Path | None:
    """
    Speichert ein Pillow-Image im temporären Verzeichnis.
    Standardmäßig wird JPEG genutzt (qual=80), um die Dateigröße zu reduzieren
    und das Einfügen ins Word-Dokument zu beschleunigen.
    """
    dir_path = Path(temp_dir) if temp_dir else TEMP_IMAGE_DIR
    dir_path.mkdir(parents=True, exist_ok=True)
    try:
        name = Path(original_path).stem
        extension = ".jpg" if format.upper() == "JPEG" else ".png"
        temp_img_path = dir_path / f"{name}{suffix}_processed{extension}"

        if format.upper() == "JPEG":
            if image.mode in ("RGBA", "LA"):
                background = Image.new("RGB", image.size, (255, 255, 255))
                background.paste(image, mask=image.split()[-1])
                image_to_save = background
            else:
                image_to_save = image.convert("RGB")
            image_to_save.save(temp_img_path, format="JPEG", quality=quality, optimize=True)
        else:
            image.save(temp_img_path, format="PNG")

        if temp_img_path.exists():
            utils.debug_log(f"Bild erfolgreich gespeichert: {temp_img_path}")
            return temp_img_path
        else:
            raise FileNotFoundError(f"Konnte Bild nicht speichern: {temp_img_path}")
    except Exception as e:
        utils.log_error(f"Fehler beim Speichern eines temporären Bildes: {e}")
        return None

def extract_text_and_images_from_pdf(
    pdf_path: Path | str,
    dpi: int = 100,
    extract_real_text: bool = True,
    progress_callback=None,
    *,
    temp_dir: Path | None = None,
) -> dict:
    """
    Rendert jede Seite als JPEG (qual=80) und extrahiert (falls vorhanden) echten PDF-Text.
    Niedriger dpi-Wert beschleunigt das Rendering und reduziert die Dateigröße.
    """
    utils.debug_log(f"Starte PDF-Verarbeitung: {pdf_path}")
    tmp_dir = Path(temp_dir) if temp_dir else TEMP_IMAGE_DIR
    tmp_dir.mkdir(parents=True, exist_ok=True)

    all_images: list[str] = []
    all_text_fragments: list[str] = []

    with fitz.open(pdf_path) as doc:
        total_pages = len(doc)
        for page_index, page in enumerate(doc, start=1):
            if progress_callback:
                progress_callback(page_index, total_pages)
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            page_image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_image_path = tmp_dir / f"{Path(pdf_path).stem}_page_{page_index:03}.jpg"

            if page_image.mode in ("RGBA", "LA"):
                background = Image.new("RGB", page_image.size, (255, 255, 255))
                background.paste(page_image, mask=page_image.split()[-1])
                page_image_rgb = background
            else:
                page_image_rgb = page_image.convert("RGB")

            page_image_rgb.save(page_image_path, format="JPEG", quality=80, optimize=True)
            all_images.append(str(page_image_path))
            utils.debug_log(
                f"[Seite {page_index}] Vollbild gerendert: {page_image_path}"
            )

            if extract_real_text:
                text_blocks = page.get_text("blocks") or []
                if text_blocks:
                    page_text = f"--- Seite {page_index} ---\n"
                    for block in text_blocks:
                        block_text = block[4].strip()
                        if block_text:
                            page_text += block_text + "\n"
                    if page_text.strip():
                        all_text_fragments.append(page_text.strip())

    final_text = "\n".join(all_text_fragments).strip()
    result = {"images": all_images, "text": final_text}
    utils.debug_log(
        f"Fertig: {pdf_path}, Seiten={len(all_images)}, Bilder={len(all_images)}, Textlänge={len(final_text)}"
    )
    return result
