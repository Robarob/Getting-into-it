from docx import Document
from docx.shared import Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from PIL import Image
from file_processing_utils import clean_text, extract_image_references
import utils
import re

def remove_invalid_xml_chars(text: str) -> str:
    pattern = re.compile(
        r'[^\u0009\u000A\u000D\u0020-\uD7FF\uE000-\uFFFD\u10000-\u10FFFF]'
    )
    return re.sub(pattern, '', text)

def write_to_word(contents, output_file, base_directory):
    """
    Schreibt extrahierten Text und Bilder in eine Word-Datei.
    """
    try:
        doc = Document()

        for file, content in contents:
            file_name = Path(file).stem
            heading_text = remove_invalid_xml_chars(file_name)
            doc.add_heading(heading_text, level=1)

            raw_text = content.get("text", "")
            if isinstance(raw_text, str) and raw_text.strip():
                cleaned_text_content = clean_text(raw_text.strip())
                cleaned_text_content = remove_invalid_xml_chars(cleaned_text_content)
                for line in cleaned_text_content.split("\n"):
                    line = line.strip()
                    if line:
                        safe_line = remove_invalid_xml_chars(line)
                        doc.add_paragraph(safe_line)
            else:
                utils.debug_log(f"Kein oder leerer Text in Datei {file}")

            images = content.get("images", [])
            for img_path in images:
                try:
                    if img_path and isinstance(img_path, (str, Path)):
                        img_file = Path(img_path)
                        if img_file.exists():
                            doc.add_picture(str(img_file), width=Cm(15))
                            utils.debug_log(f"Bild erfolgreich hinzugefügt: {img_file}")
                        else:
                            utils.log_error(f"Bilddatei existiert nicht: {img_path}")
                    else:
                        utils.log_error(f"Ungültiger oder leerer Bildpfad erkannt: {img_path}")
                except Exception as img_error:
                    utils.log_error(f"Fehler beim Einfügen eines Bildes {img_path}: {img_error}")

        output_file_path = Path(output_file)
        output_file_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_file_path)
        utils.debug_log(f"Word-Datei erfolgreich gespeichert: {output_file_path}")

    except Exception as e:
        utils.log_error(f"Fehler beim Schreiben in die Word-Datei {output_file}: {e}")
        raise

def add_bookmark(paragraph, bookmark_name):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.insert(0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "0")
    paragraph._p.append(end)

def add_hyperlink(paragraph, bookmark_name, text, color=RGBColor(0, 0, 255), underline=True):
    hyperlink = paragraph.add_run(text)
    hyperlink.font.color.rgb = color
    if underline:
        hyperlink.font.underline = True
